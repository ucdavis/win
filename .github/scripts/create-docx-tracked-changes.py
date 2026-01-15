#!/usr/bin/env python3
"""
Script to compare DOCX files and create a version with tracked changes.
This compares the PR's DOCX files with the published versions from gh-pages.
"""

import os
import sys
import subprocess
from pathlib import Path

def checkout_base_docx(base_ref='origin/gh-pages', target_dir='/tmp/base-docx'):
    """Check out the base DOCX files from gh-pages for comparison."""
    target_path = Path(target_dir)
    target_path.mkdir(parents=True, exist_ok=True)
    
    try:
        subprocess.run(['git', 'fetch', 'origin', 'gh-pages:gh-pages'], 
                      check=False, capture_output=True)
        
        result = subprocess.run(
            ['git', 'ls-tree', '-r', '--name-only', base_ref],
            capture_output=True,
            text=True,
            check=False
        )
        
        if result.returncode == 0:
            files = [f for f in result.stdout.split('\n') if f.endswith('.docx')]
            
            for file in files:
                output_path = target_path / file
                output_path.parent.mkdir(parents=True, exist_ok=True)
                
                with open(output_path, 'wb') as f:
                    subprocess.run(
                        ['git', 'show', f'{base_ref}:{file}'],
                        stdout=f,
                        check=False
                    )
            
            return target_path if files else None
        
        return None
    except Exception as e:
        print(f"Could not check out base DOCX: {e}", file=sys.stderr)
        return None

def create_docx_with_tracked_changes(old_docx_path, new_docx_path, output_path):
    """
    Create a DOCX file with tracked changes showing differences.
    This uses python-docx to enable track changes and show revisions.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import difflib
        import shutil
        
        # Copy the new document to the output path
        shutil.copy2(new_docx_path, output_path)
        
        # Load the output document
        output_doc = Document(output_path)
        
        # Enable track changes in the document settings
        settings = output_doc.settings
        settings_element = settings.element
        
        # Add trackRevisions element if it doesn't exist
        track_revisions = settings_element.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = OxmlElement('w:trackRevisions')
            settings_element.append(track_revisions)
        
        # Load old and new documents for comparison
        old_doc = Document(old_docx_path)
        new_doc = Document(new_docx_path)
        
        # Get paragraphs from both documents
        old_paragraphs = [p.text for p in old_doc.paragraphs]
        new_paragraphs = [p.text for p in new_doc.paragraphs]
        
        # Use difflib to find differences at paragraph level
        matcher = difflib.SequenceMatcher(None, old_paragraphs, new_paragraphs)
        has_changes = False
        
        # Process each operation
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                has_changes = True
                for idx in range(j1, j2):
                    if idx < len(output_doc.paragraphs):
                        para = output_doc.paragraphs[idx]
                        for run in para.runs:
                            ins = OxmlElement('w:ins')
                            ins.set(qn('w:id'), str(idx))
                            ins.set(qn('w:author'), 'PR Preview')
                            ins.set(qn('w:date'), '2024-01-01T00:00:00Z')
                            
                            run_element = run._element
                            parent = run_element.getparent()
                            parent.insert(parent.index(run_element), ins)
                            parent.remove(run_element)
                            ins.append(run_element)
                            
            elif tag == 'insert':
                has_changes = True
                for idx in range(j1, j2):
                    if idx < len(output_doc.paragraphs):
                        para = output_doc.paragraphs[idx]
                        for run in para.runs:
                            ins = OxmlElement('w:ins')
                            ins.set(qn('w:id'), str(1000 + idx))
                            ins.set(qn('w:author'), 'PR Preview')
                            ins.set(qn('w:date'), '2024-01-01T00:00:00Z')
                            
                            run_element = run._element
                            parent = run_element.getparent()
                            parent.insert(parent.index(run_element), ins)
                            parent.remove(run_element)
                            ins.append(run_element)
        
        # Save the document with tracked changes enabled
        output_doc.save(output_path)
        
        if has_changes:
            print(f"  ✓ Created DOCX with tracked changes: {output_path}")
        else:
            print(f"  ✓ Created DOCX with track changes enabled (no paragraph-level changes found): {output_path}")
        
        return True
            
    except ImportError:
        print("  ⚠ Warning: python-docx not available")
        print("    Copying new DOCX without tracked changes markup")
        import shutil
        shutil.copy2(new_docx_path, output_path)
        return True
    except Exception as e:
        print(f"  ✗ Error creating tracked changes DOCX: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        try:
            import shutil
            shutil.copy2(new_docx_path, output_path)
            print(f"  ✓ Copied DOCX without tracked changes as fallback")
            return True
        except:
            return False

def process_docx_file(new_docx_path, base_docx_dir):
    """Process a single DOCX file: fetch old version, compare, and create tracked changes version."""
    print(f"Processing {new_docx_path}...")
    
    if not base_docx_dir:
        print("  No base DOCX directory available")
        return
    
    new_path = Path(new_docx_path)
    relative_path = new_path.name
    
    # Try different base paths
    base_path = Path(base_docx_dir) / relative_path
    if not base_path.exists():
        # Try in _site/chapters/
        base_path = Path(base_docx_dir) / "_site" / "chapters" / relative_path
    if not base_path.exists():
        # Try in chapters/
        base_path = Path(base_docx_dir) / "chapters" / relative_path
    
    if not base_path.exists():
        print(f"  Base DOCX not found: {base_path}")
        return
    
    # Create output path with tracked changes
    output_path = new_path.parent / f"{new_path.stem}-tracked-changes.docx"
    
    print(f"  Output will be: {output_path}")
    
    success = create_docx_with_tracked_changes(base_path, new_path, output_path)
    
    if success:
        print(f"  Successfully created: {output_path}")

def main():
    docx_dir = os.getenv('DOCX_DIR', './_site')
    
    print("="*60)
    print("DOCX Tracked Changes Creation")
    print("="*60)
    
    print("\n1. Checking out base DOCX files from gh-pages...")
    base_docx_dir = checkout_base_docx()
    
    if not base_docx_dir:
        print("⚠ Warning: Could not check out base DOCX files")
        print("   (This is normal for:")
        print("    - First PR to a new repository")
        print("    - If gh-pages branch doesn't have DOCX files yet)")
        print("   Skipping DOCX tracked changes creation.")
        return
    else:
        print(f"✓ Base DOCX checked out to {base_docx_dir}")
    
    # Find all DOCX files in the chapters directory
    docx_files = list(Path(docx_dir).glob("chapters/*-handout.pdf"))
    
    # Also look for regular DOCX files if any exist
    docx_files_direct = list(Path(docx_dir).glob("chapters/*.docx"))
    if docx_files_direct:
        docx_files.extend(docx_files_direct)
    
    if not docx_files:
        print("\n⚠ No DOCX files found in output directory")
        return
    
    print(f"\n2. Found {len(docx_files)} DOCX file(s) to process:")
    for docx_file in docx_files:
        print(f"   - {docx_file.name}")
    
    print("\n3. Creating tracked changes versions:")
    for docx_file in docx_files:
        process_docx_file(docx_file, base_docx_dir)
    
    print("\n" + "="*60)
    print("DOCX processing complete")
    print("="*60)

if __name__ == '__main__':
    main()
